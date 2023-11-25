from docx import Document
from docx.shared import Inches


def replace_string_paragraph(filename, current_word, replacement_word):
    document = Document(filename)
    for paragraph in document.paragraphs:
        if current_word in paragraph.text:
            paragraph.text = replacement_word
    document.save(filename)


def replace_string_table(filename, current_word, replacement_word):
    document = Document(filename)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if current_word in p.text:
                        inline = paragraph.runs
                        for i in range(len(inline)):
                            if current_word in inline[i].text:
                                text = inline[i].text.replace(current_word, replacement_word)
                                inline[i].text = text
    document.save(filename)

